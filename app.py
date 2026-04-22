# app.py — Hackathon PDF Copilot Full Version with Download Options
import os
import re
import io
import json
import tempfile
from typing import List, Dict, Tuple

import streamlit as st
from dotenv import load_dotenv

import pdfplumber
from fpdf import FPDF
from docx import Document
from PIL import Image
from wordcloud import WordCloud

# Optional OCR
try:
    from pdf2image import convert_from_bytes
    import pytesseract
    OCR_AVAILABLE = True
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
except:
    OCR_AVAILABLE = False

# Retrieval for Q&A
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Gemini AI
import google.generativeai as genai

# Named Entity Recognition
import spacy
nlp = spacy.load("en_core_web_sm")

# -------------------------
# Page config + CSS styling
# -------------------------
st.set_page_config(page_title="✨ APEX SUMMARIZER CO-PILOT", layout="wide", page_icon="📄")

st.markdown("""
<style>
.stApp { background: linear-gradient(135deg,#1e1e2f,#3a506b,#5bc0be,#6fffe9); background-size:400% 400%; color:#f0f0f0; font-family:'Segoe UI',sans-serif; }
.card { background: rgba(0,0,0,0.4); border-radius:15px; padding:20px; backdrop-filter:blur(10px); margin-bottom:20px; color:#f0f0f0; }
.stButton>button { background-color:#4b9ce2; color:white; border-radius:10px; padding:0.6em 1.2em; font-weight:bold; border:none; transition:0.3s; }
.stButton>button:hover { background-color:#1b6ec2; transform:scale(1.05);}
.stDownloadButton>button { background-color:#0b3558; color:white; border-radius:8px; padding:0.5rem 0.9rem;}
.term-chip { display:inline-block; margin:4px 6px 6px 0; background:#ffd966; color:#0b2030; font-weight:700; padding:4px 8px; border-radius:10px;}
.codebox { background: rgba(255,255,255,0.1); padding:10px; border-radius:8px; border:1px solid rgba(255,255,255,0.2);}
</style>
""", unsafe_allow_html=True)

st.markdown("<h1 style='text-align:center'>✨ APEX SUMMARIZER CO-PILOT</h1>", unsafe_allow_html=True)
st.caption("Upload PDF → Get summary, section insights, glossary, word cloud, named entities, Q&A, and export options.")

# -------------------------
# Load Gemini API Key
# -------------------------
load_dotenv()
API_KEY = os.getenv("GEMINI_API_KEY", "").strip()
if not API_KEY:
    st.error("❌ GEMINI_API_KEY not found. Create a `.env` with `GEMINI_API_KEY=...`.")
    st.stop()
genai.configure(api_key=API_KEY)
MODEL_ID = "gemini-2.0-flash"

# -------------------------
# PDF Unicode-safe class
# -------------------------
class PDFUnicode(FPDF):
    def __init__(self, orientation='P', unit='mm', format='A4'):
        super().__init__(orientation, unit, format)
        font_path = os.path.join(os.path.dirname(__file__), "DejaVuSans.ttf")
        if os.path.exists(font_path):
            self.add_font("DejaVu", "", font_path, uni=True)
            self.add_font("DejaVu", "B", font_path, uni=True)
        self.set_auto_page_break(auto=True, margin=15)

# -------------------------
# Helper functions
# -------------------------
def safe_text(text, max_len=80):
    if not text: return ""
    words=[]
    for w in text.split():
        if len(w)>max_len:
            words.extend([w[i:i+max_len] for i in range(0,len(w),max_len)])
        else: words.append(w)
    return " ".join(words)

def extract_text_pdfplumber(pdf_bytes: bytes) -> Tuple[str,int]:
    txt_parts=[]; pages=0
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(pdf_bytes); tmp_path=tmp.name
    try:
        with pdfplumber.open(tmp_path) as pdf:
            pages=len(pdf.pages)
            for p in pdf.pages: txt_parts.append(p.extract_text() or "")
    except: return "",0
    return "\n".join(txt_parts).strip(), pages

def extract_text_ocr(pdf_bytes: bytes) -> str:
    if not OCR_AVAILABLE: return ""
    try:
        images = convert_from_bytes(pdf_bytes, dpi=200)
        lines=[]
        for img in images: lines.append(pytesseract.image_to_string(img, lang="eng"))
        return "\n".join(lines).strip()
    except: return ""

def chunk_text(s: str, max_chars: int=9000, overlap: int=500) -> List[str]:
    s=s.strip(); chunks=[]; i=0; n=len(s)
    while i<n:
        j=min(i+max_chars,n)
        part=s[i:j]; last_stop=part.rfind(".")
        if last_stop>1500 and j!=n: j=i+last_stop+1; part=s[i:j]
        chunks.append(part); i=max(j-overlap,j)
    return chunks

def gemini_call(prompt: str) -> str:
    try:
        model = genai.GenerativeModel(MODEL_ID)
        resp = model.generate_content(prompt)
        if hasattr(resp,"text") and resp.text: return resp.text.strip()
        return resp.candidates[0].content.parts[0].text.strip()
    except: return ""

# -------------------------
# Summarization / Glossary / Sections
# -------------------------
def detect_language(text: str) -> str:
    lang = gemini_call(f"Detect language of this text. Reply with language name only:\n{text[:4000]}")
    return (lang or "English").splitlines()[0].strip()

def summarize_document(text:str,length="short",tone="neutral",lang="auto") -> str:
    base_lang=detect_language(text)
    target_lang=base_lang if lang=="auto" else lang
    chunks=chunk_text(text)
    partial=[]
    for ch in chunks:
        s=gemini_call(f"Summarize in 3-4 sentences, tone {tone}, in {target_lang}:\n{ch}")
        if s: partial.append(s)
    if not partial: return ""
    final=gemini_call(f"Combine partial summaries in 3-4 sentences, tone {tone}:\n{''.join(partial)}")
    return final or ""

def section_summaries(text:str,lang="auto") -> Dict[str,str]:
    base_lang=detect_language(text); target_lang=base_lang if lang=="auto" else lang
    prompt=(f"Split document into sections with 1-2 sentence summaries. Return STRICT JSON [{{\"section\":\"...\",\"summary\":\"...\"}}].\n\nText:\n{text[:15000]}")
    raw=gemini_call(prompt); out={}
    try:
        m=re.search(r"\[.*\]",raw,re.S); js=json.loads(m.group(0) if m else raw)
        for it in js: out[it.get("section","Section")]=it.get("summary","")
    except: out={"Introduction":text[:500],"Body":text[:1000],"Conclusion":text[-500:]}
    return out

def build_glossary(text:str,max_terms=12,lang="auto") -> List[Dict[str,str]]:
    base_lang=detect_language(text); target_lang=base_lang if lang=="auto" else lang
    prompt=f"Extract {max_terms} terms with explanations as JSON [{{\"term\":\"\",\"definition\":\"\"}}] in {target_lang}:\n{text[:15000]}"
    raw=gemini_call(prompt); items=[]
    try:
        m=re.search(r"\[.*\]",raw,re.S); js=json.loads(m.group(0) if m else raw)
        for it in js: items.append({"Term":it.get("term",""),"Explanation":it.get("definition","")})
    except:
        for line in raw.splitlines():
            if " - " in line:
                t,d=line.split(" - ",1); items.append({"Term":t.strip(),"Explanation":d.strip()})
    seen,dedup=set(),[]
    for it in items:
        k=it["Term"].lower()
        if k not in seen: seen.add(k); dedup.append(it)
    return dedup[:max_terms]

def build_wordcloud_image(text:str):
    wc = WordCloud(width=1000,height=400,background_color=None,mode="RGBA",collocations=False,max_words=150)
    return wc.generate(text).to_image()

# -------------------------
# Q&A Retriever
# -------------------------
def make_retriever(chunks:List[str]):
    vec=TfidfVectorizer(stop_words="english"); X=vec.fit_transform(chunks)
    return vec,X

def retrieve_relevant(chunks,vec,X,query,k=5):
    qv=vec.transform([query]); sims=cosine_similarity(qv,X).ravel()
    return [chunks[i] for i in sims.argsort()[::-1][:k] if sims[i]>0]

def qa_answer(question:str,retrieved:List[str],lang="auto") -> str:
    if not retrieved: return "No relevant context found."
    ctx="\n\n---\n\n".join(retrieved[:5])
    prompt=f"Answer the question based ONLY on context. Write in {lang}.\nContext:\n{ctx}\nQuestion:\n{question}"
    return gemini_call(prompt)

# -------------------------
# Named Entity Recognition
# -------------------------
def extract_entities(text:str) -> List[Dict[str,str]]:
    doc = nlp(text)
    entities=[]
    for ent in doc.ents:
        entities.append({"text":ent.text,"label":ent.label_})
    return entities

# -------------------------
# Export functions
# -------------------------
def to_pdf(summary,sections,glossary,entities) -> bytes:
    pdf=PDFUnicode(); pdf.add_page()
    pdf.set_font("DejaVu","B",16); pdf.cell(0,10,"AI PDF Report",ln=True)
    pdf.set_font("DejaVu","",12); pdf.multi_cell(0,7,safe_text("Summary:\n"+(summary or "—")))
    if sections:
        pdf.ln(2); pdf.set_font("DejaVu","B",14); pdf.cell(0,10,"Section Summaries",ln=True)
        for k,v in sections.items():
            pdf.set_font("DejaVu","B",12); pdf.multi_cell(0,7,safe_text(k))
            pdf.set_font("DejaVu","",12); pdf.multi_cell(0,7,safe_text(v)); pdf.ln(1)
    if glossary:
        pdf.ln(2); pdf.set_font("DejaVu","B",14); pdf.cell(0,10,"Glossary",ln=True)
        pdf.set_font("DejaVu","",12)
        for it in glossary:
            pdf.set_font("DejaVu","B",12); pdf.multi_cell(0,7,safe_text(it["Term"]))
            pdf.set_font("DejaVu","",12); pdf.multi_cell(0,7,safe_text(it["Explanation"])); pdf.ln(1)
    if entities:
        pdf.ln(2); pdf.set_font("DejaVu","B",14); pdf.cell(0,10,"Named Entities",ln=True)
        pdf.set_font("DejaVu","",12)
        for ent in entities:
            pdf.multi_cell(0,7,f"{ent['text']} ({ent['label']})")
    out=pdf.output(dest="S")
    if isinstance(out,bytearray): out=bytes(out)
    elif isinstance(out,str): out=out.encode("utf-8","ignore")
    return out

def to_docx(summary,sections,glossary,entities) -> bytes:
    doc=Document(); doc.add_heading("AI PDF Report",level=1)
    doc.add_heading("Summary",level=2); doc.add_paragraph(summary or "—")
    if sections:
        doc.add_heading("Section Summaries",level=2)
        for k,v in sections.items(): doc.add_heading(k,level=3); doc.add_paragraph(v)
    if glossary:
        doc.add_heading("Glossary",level=2)
        for it in glossary: doc.add_paragraph(f"{it['Term']}: {it['Explanation']}")
    if entities:
        doc.add_heading("Named Entities",level=2)
        for ent in entities: doc.add_paragraph(f"{ent['text']} ({ent['label']})")
    bio=io.BytesIO(); doc.save(bio); bio.seek(0); return bio.read()

def to_markdown(summary,sections,glossary,entities) -> str:
    lines=["# AI PDF Report","", "## Summary", summary or "—"]
    if sections:
        lines+=["","## Section Summaries"]
        for k,v in sections.items(): lines+=["### "+k,v,""]
    if glossary:
        lines+=["","## Glossary"]
        for it in glossary: lines+=[f"- **{it['Term']}** — {it['Explanation']}"]
    if entities:
        lines+=["","## Named Entities"]
        for ent in entities: lines+=[f"- **{ent['text']}** ({ent['label']})"]
    return "\n".join(lines)

# -------------------------
# Sidebar and main processing
# -------------------------
with st.sidebar:
    st.markdown("### ⚙️ Settings")
    summary_len=st.selectbox("Summary length",["very short","short","medium","detailed"],index=1)
    tone=st.selectbox("Tone",["neutral","formal","friendly","executive"],index=0)
    out_lang=st.selectbox("Output language",["auto","English","Hindi","Tamil","Telugu","Spanish","French","Arabic","German","Chinese"],index=0)
    max_terms=st.slider("Max glossary terms",5,30,12,1)
    do_sections=st.checkbox("Generate section-wise summaries",value=True)
    make_wordcloud=st.checkbox("Show keyword word cloud",value=True)
    allow_ocr=st.checkbox("Use OCR if PDF has no text",value=True)
    st.markdown("---")
    st.caption("Tip: For OCR install Tesseract and set path if needed.")

uploaded = st.file_uploader("📤 Upload a PDF", type=["pdf"])
if uploaded:
    file_bytes = uploaded.read()
    text,pages = extract_text_pdfplumber(file_bytes)
    if not text and allow_ocr: text = extract_text_ocr(file_bytes)
    if not text: st.error("No text found."); st.stop()
    
    st.markdown(f"<div class='card'>Pages: <b>{pages}</b> | Characters: <b>{len(text)}</b></div>",unsafe_allow_html=True)
    with st.expander("👀 Preview extracted text (first 3000 chars)"):
        st.code(text[:3000]+("…" if len(text)>3000 else ""), language="markdown")
    
    # Buttons
    c1,c2,c3 = st.columns(3)
    run_all = c1.button("✨ Run All")
    just_summary = c2.button("📝 Summary only")
    just_glossary = c3.button("📚 Glossary only")
    
    # Session state
    if "summary" not in st.session_state: st.session_state.summary=""
    if "sections" not in st.session_state: st.session_state.sections={}
    if "glossary" not in st.session_state: st.session_state.glossary=[]
    if "chunks" not in st.session_state: st.session_state.chunks=chunk_text(text)
    if "entities" not in st.session_state: st.session_state.entities=[]
    
    # Run processing
    if run_all or just_summary: st.session_state.summary = summarize_document(text,length=summary_len,tone=tone,lang=out_lang)
    if run_all or just_glossary: st.session_state.glossary = build_glossary(text,max_terms=max_terms,lang=out_lang)
    if run_all and do_sections: st.session_state.sections = section_summaries(text,lang=out_lang)
    if run_all: st.session_state.entities = extract_entities(text)
    
    # Display outputs
    if st.session_state.summary: st.markdown("<div class='card'><h3>📝 Summary</h3></div>",unsafe_allow_html=True); st.write(st.session_state.summary)
    if st.session_state.sections: st.markdown("<div class='card'><h3>📑 Section Summaries</h3></div>",unsafe_allow_html=True)
    for sec,summ in st.session_state.sections.items(): st.markdown(f"**{sec}**"); st.write(summ)
    if st.session_state.glossary: st.markdown("<div class='card'><h3>📚 Glossary</h3></div>",unsafe_allow_html=True)
    chips = " ".join([f"<span class='term-chip'>{it['Term']}</span>" for it in st.session_state.glossary])
    st.markdown(chips,unsafe_allow_html=True)
    for it in st.session_state.glossary: st.markdown(f"- **{it['Term']}** — {it['Explanation']}")
    if make_wordcloud:
        try: img = build_wordcloud_image(text); st.markdown("<div class='card'><h3>☁️ Keyword Cloud</h3></div>",unsafe_allow_html=True); st.image(img,use_container_width=True)
        except: st.warning("Word cloud unavailable")
    if st.session_state.entities: st.markdown("<div class='card'><h3>🧩 Named Entities</h3></div>",unsafe_allow_html=True)
    for e in st.session_state.entities: st.markdown(f"- **{e['text']}** ({e['label']})")
    
    # Q&A Section
st.markdown("<div class='card'><h3>💬 Ask Questions</h3></div>", unsafe_allow_html=True)

# Text input for user question
qa_q = st.text_input("Type your question here:")

# Only run if user typed something
if qa_q:
    if "vec" not in st.session_state:
        st.session_state.vec, st.session_state.X = make_retriever(st.session_state.chunks)
    
    # Retrieve relevant chunks and generate answer
    rel = retrieve_relevant(st.session_state.chunks, st.session_state.vec, st.session_state.X, qa_q, k=6)
    ans = qa_answer(qa_q, rel, lang=out_lang)
    
    # Display the answer
    st.markdown("**Answer:**")
    st.write(ans)
    
    # Optional: show the supporting context
    with st.expander("Supporting context"):
        for i, c in enumerate(rel, 1):
            st.markdown(
                f"<div class='codebox'><b>Chunk {i}</b><br>{c[:1200]}{'…' if len(c)>1200 else ''}</div>",
                unsafe_allow_html=True
            )
